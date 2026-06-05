"""
Generates USER_GUIDE.docx — a short, plain-language manual for non-technical
users (managers). Styled to match the app: Calibri + cobalt-blue headers.
Run:  python make_user_guide.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── brand (matches backend/formatters/compact_ats.py) ────────────────────────
FONT = "Calibri"
COBALT = RGBColor(0, 71, 171)
CLAY = RGBColor(0xC9, 0x64, 0x42)
INK = RGBColor(0x20, 0x20, 0x1E)
GRAY = RGBColor(0x5A, 0x5A, 0x5A)
BOX_BG = "EAF1FB"   # pale cobalt
WARN_BG = "FBF0EC"  # pale clay


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def run(p, text, *, size=10.5, bold=False, color=INK, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return r


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run(p, text, size=13, bold=True, color=COBALT)
    return p


def callout(doc, title, lines, bg, accent):
    """A single-cell shaded box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(t)
    cell = t.cell(0, 0)
    _shade(cell, bg)
    _set_cell_margins(cell, top=120, bottom=120)
    cell.paragraphs[0].text = ""
    ph = cell.paragraphs[0]
    ph.paragraph_format.space_after = Pt(2)
    run(ph, title, size=10.5, bold=True, color=accent)
    for ln in lines:
        pl = cell.add_paragraph()
        pl.paragraph_format.space_after = Pt(1)
        run(pl, ln, size=10, color=INK)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)
    sp.paragraph_format.space_before = Pt(0)
    return t


def step(doc, n, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run(p, f"{n}  ", size=12, bold=True, color=CLAY)
    run(p, title, size=12, bold=True, color=INK)
    b = doc.add_paragraph()
    b.paragraph_format.left_indent = Inches(0.28)
    b.paragraph_format.space_after = Pt(4)
    # body is a list of (text, bold) segments OR a plain string
    if isinstance(body, str):
        run(b, body, size=10.5, color=GRAY)
    else:
        for seg, is_bold in body:
            run(b, seg, size=10.5, bold=is_bold, color=GRAY if not is_bold else INK)
    return p


# ── build ────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.55)
sec.bottom_margin = Inches(0.45)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)

# Title
title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(0)
run(title, "Resume Formatter", size=24, bold=True, color=COBALT)
sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(2)
run(sub, "A 1-minute guide to turning any resume into a clean, ready-to-send file.",
    size=11.5, italic=True, color=GRAY)

intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(8)
run(intro, "Upload a resume, review it on screen, and download a polished Word file. ",
    size=10.5, color=INK)
run(intro, "Nothing is ever added, nothing is ever lost", size=10.5, bold=True, color=INK)
run(intro, " — the result only ever contains words from the original resume.", size=10.5, color=INK)

# Open the app
callout(
    doc,
    "Before you start  —  open the app",
    [
        "Open the link your team shared:  ______________________________________",
        "Nothing to install, no login. Works in any browser (Chrome, Edge, Safari).",
    ],
    BOX_BG, COBALT,
)

# Steps
heading(doc, "The 4 steps")

step(doc, "1", "Upload your resume",
     [("On the home screen click ", False), ("Format your resume", True),
      (". Drag a PDF or Word file into the box (or click ", False), ("Paste text", True),
      (" to paste it in). Then click ", False), ("Format resume", True),
      (" and wait a few seconds.", False)])

step(doc, "2", "Review it on the right",
     [("Your resume appears as a real page. On the left, every section has three buttons — ", False),
      ("Keep", True), (" it as-is, ", False), ("Skip", True),
      (" to leave it out, or ", False), ("Edit", True),
      (" to change the wording. The page on the right updates as you go.", False)])

# extras under step 2
ex = doc.add_paragraph()
ex.paragraph_format.left_indent = Inches(0.28)
ex.paragraph_format.space_after = Pt(5)
run(ex, "Handy extras:  ", size=10, bold=True, color=COBALT)
run(ex, "drag the ⠿ handle to reorder sections · turn GPA / coursework on or off · "
        "use “Edit name, contact & headline” at the top to fix the header.",
    size=10, color=GRAY)

step(doc, "3", "Approve everything",
     [("Click ", False), ("✓ Approve all remaining", True),
      (" to accept every section in one click (or review them one by one). "
       "The download button unlocks once the progress bar turns green.", False)])

step(doc, "4", "Download the file",
     [("Click ", False), ("↓ Download DOCX", True),
      (". The clean resume saves to your ", False), ("Downloads", True),
      (" folder — open it in Word, print it, or email it. Done.", False)])

# Cheat sheet
heading(doc, "Button cheat-sheet")
rows = [
    ("Keep", "Include this section exactly as it is."),
    ("Skip", "Leave this section out of the final file."),
    ("Edit", "Reword the text (you can only edit words already there)."),
    ("⠿  drag", "Move a section up or down."),
    ("Approve all remaining", "Accept everything at once — fastest way to finish."),
    ("Download DOCX", "Save the finished resume as a Word file."),
    ("Format another resume", "Start over with a new file."),
]
tbl = doc.add_table(rows=len(rows), cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
_no_borders(tbl)
tbl.columns[0].width = Inches(1.9)
tbl.columns[1].width = Inches(4.6)
for i, (btn, desc) in enumerate(rows):
    c0, c1 = tbl.rows[i].cells
    c0.width = Inches(1.9)
    c1.width = Inches(4.6)
    if i % 2 == 0:
        _shade(c0, "F4F7FC")
        _shade(c1, "F4F7FC")
    _set_cell_margins(c0, top=40, bottom=40)
    _set_cell_margins(c1, top=40, bottom=40)
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    run(p0, btn, size=10, bold=True, color=CLAY)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    run(p1, desc, size=10, color=INK)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Troubleshooting
callout(
    doc,
    "If something goes wrong",
    [
        "“Couldn’t reach the server” — it’s just waking up (the first use after a quiet "
        "spell takes ~30–60 sec). Wait a moment and click Download again; your edits are safe.",
        "Want to pause? Click “Download my edits” to save your work and come back later.",
        "The tool never invents or deletes content — if a line looks misplaced, just Edit or "
        "drag it. Every word from the original is always kept.",
    ],
    WARN_BG, CLAY,
)

# Footer line
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.paragraph_format.space_before = Pt(6)
run(foot, "Resume Formatter  ·  the cleanest way to format a resume — every word preserved.",
    size=9, italic=True, color=GRAY)

out = "USER_GUIDE.docx"
doc.save(out)
print("wrote", out)
