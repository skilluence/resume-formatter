from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Response
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Any, Dict
from dotenv import load_dotenv
import logging
import os
import re
import shutil
import tempfile
import uuid

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("resume.api")

app = FastAPI(title="Resume Formatter API")

# Comma-separated origins, or "*" for any. Logged at startup so a CORS/"failed to
# fetch" misconfig is visible in the deploy logs.
ALLOWED_ORIGINS = [o.strip() for o in os.getenv("ALLOWED_ORIGINS", "*").split(",") if o.strip()] or ["*"]
logger.info("[cors] allowed origins: %s", ALLOWED_ORIGINS)

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Uploads are parsed then deleted immediately — only a transient scratch dir.
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "..", "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _safe_filename(name: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", (name or "").strip()).strip("_")
    return safe or "Resume"


@app.get("/")
def root():
    """Health check / warm-up ping (the frontend hits this to wake a cold server)."""
    return {"status": "Resume Formatter API is running"}


# ─────────────────────────────── parse / structure ──────────────────────────

async def _read_resume_text(file, plain_text) -> str:
    """Turn an uploaded file (PDF/DOCX) or pasted text into raw resume text.

    Shared by /format and /tailor. Raises a clear 400 (never a 500) on every
    failure path, with pasting text always available as the fallback."""
    if not file and not plain_text:
        raise HTTPException(status_code=400, detail="Provide a file or paste your resume text.")

    if not file:
        raw_text = plain_text or ""
    else:
        ext = os.path.splitext(file.filename or "")[1].lower()
        if ext == ".doc":
            raise HTTPException(
                status_code=400,
                detail="Old .doc files aren't supported — please save it as .docx or PDF, or paste the text.",
            )
        if ext not in (".pdf", ".docx"):
            raise HTTPException(
                status_code=400,
                detail="Only PDF or DOCX files are supported. You can also paste the text instead.",
            )
        tmp_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4().hex[:8]}{ext}")
        try:
            with open(tmp_path, "wb") as f:
                f.write(await file.read())
            from parsers.pdf_parser import extract_text_from_pdf
            from parsers.docx_parser import extract_text_from_docx
            try:
                raw_text = (
                    extract_text_from_pdf(tmp_path) if ext == ".pdf"
                    else extract_text_from_docx(tmp_path)
                )
            except Exception:
                logger.exception("[parse] failed to parse %s", file.filename)
                raise HTTPException(
                    status_code=400,
                    detail="We couldn't read this file — it may be a scanned image, password-protected, "
                           "or corrupted. Try a text-based PDF/DOCX, or paste the text.",
                )
        finally:
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    if not (raw_text or "").strip():
        raise HTTPException(
            status_code=400,
            detail="No text could be read from this resume — if it's a scanned image, paste the text instead.",
        )
    return raw_text


@app.post("/format")
async def format_resume(file: UploadFile = File(None), plain_text: str = Form(None)):
    """Parse an uploaded resume (or pasted text) into structured JSON for review.

    No document is written here — the user reviews/edits/approves first, then
    calls /build. Every failure path returns a clear 400 with guidance instead
    of a 500, and pasting text is always available as a fallback."""
    raw_text = await _read_resume_text(file, plain_text)
    from structurer import structure_resume
    resume = structure_resume(raw_text)
    return JSONResponse({"candidate_name": resume.get("name") or "Unknown", "resume": resume})


# ─────────────────────────────── build / download ───────────────────────────

class BuildRequest(BaseModel):
    resume: Dict[str, Any]


class TailorBuildRequest(BaseModel):
    kind: str  # "resume" | "cover_letter" | "email"
    format: str = "docx"  # "docx" | "pdf"
    tailored_resume: Dict[str, Any] = {}
    cover_letter: Dict[str, Any] = {}
    email: Dict[str, Any] = {}


def _require_named_resume(resume):
    if not isinstance(resume, dict) or not (resume.get("name") or "").strip():
        raise HTTPException(status_code=400, detail="A resume with at least a name is required.")


def _fallback_docx(resume: dict, path: str):
    """Last-resort minimal DOCX: the name plus every field as plain text. Even if
    the rich formatter somehow fails, the user still gets a downloadable file with
    nothing lost."""
    from docx import Document
    doc = Document()
    doc.add_paragraph((resume.get("name") or "Resume").strip())

    def emit(value, label=""):
        if isinstance(value, str):
            if value.strip():
                doc.add_paragraph(f"{label}{value.strip()}")
        elif isinstance(value, dict):
            for k, v in value.items():
                emit(v, f"{k}: " if isinstance(v, str) else "")
        elif isinstance(value, list):
            for v in value:
                emit(v)

    for key in ("headline", "contact", "summary", "skills", "experience",
                "projects", "certifications", "additional_sections", "education"):
        emit(resume.get(key))
    doc.save(path)


def _render_docx(resume: dict, path: str):
    from formatters.compact_ats import format_compact
    try:
        format_compact(resume, path)
    except Exception:
        logger.exception("[build] rich formatter failed; using text fallback")
        _fallback_docx(resume, path)


@app.post("/build")
def build_docx(req: BuildRequest):
    """Render the approved resume JSON and stream the DOCX back in ONE request.

    No disk persistence and no second download request, so there's nothing to
    404 if the server restarts/cold-starts between calls."""
    _require_named_resume(req.resume)
    filename = f"{_safe_filename(req.resume.get('name'))}.docx"
    tmpdir = tempfile.mkdtemp(prefix="resume-")
    try:
        path = os.path.join(tmpdir, filename)
        _render_docx(req.resume, path)
        with open(path, "rb") as f:
            data = f.read()
    except Exception as e:
        logger.exception("[build] could not generate DOCX")
        raise HTTPException(status_code=500, detail=f"Couldn't generate the DOCX: {e}")
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
    return Response(
        content=data,
        media_type=DOCX_MEDIA_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _convert_to_pdf(docx_path: str, pdf_path: str, workdir: str):
    """DOCX → PDF. Windows drives Word via docx2pdf; Linux/Mac use LibreOffice
    headless with an isolated profile (concurrency-safe). Raises on failure."""
    if os.name == "nt":
        import subprocess
        import sys
        # MS Word over COM only works on a process's STA main thread; FastAPI runs
        # sync endpoints in a worker thread, so an inline convert() raises
        # 0x800706B5 ("the interface is unknown"). Run a dedicated helper script in
        # a FRESH subprocess: clean main-thread COM apartment + a private Word
        # instance (DispatchEx) that ignores any stuck/zombie Word.
        helper = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_win_pdf_convert.py")
        result = subprocess.run(
            [sys.executable, helper, docx_path, pdf_path],
            capture_output=True,
            text=True,
            timeout=180,
        )
        if not os.path.exists(pdf_path):
            err = (result.stderr or result.stdout or "").strip()[:400]
            raise RuntimeError(f"Word PDF export failed (exit={result.returncode}); {err}")
    else:
        import subprocess
        result = subprocess.run(
            [
                "soffice",
                f"-env:UserInstallation=file://{os.path.join(workdir, 'lo-profile')}",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(pdf_path),
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0 or not os.path.exists(pdf_path):
            raise RuntimeError(f"soffice exit={result.returncode}; stderr={result.stderr.strip()[:400]}")


@app.post("/build/pdf")
def build_pdf(req: BuildRequest):
    """Render + convert to PDF and stream it back in one request."""
    _require_named_resume(req.resume)
    safe = _safe_filename(req.resume.get("name"))
    tmpdir = tempfile.mkdtemp(prefix="resume-")
    try:
        docx_path = os.path.join(tmpdir, f"{safe}.docx")
        pdf_path = os.path.join(tmpdir, f"{safe}.pdf")
        _render_docx(req.resume, docx_path)
        try:
            _convert_to_pdf(docx_path, pdf_path, tmpdir)
        except FileNotFoundError:
            logger.exception("[build/pdf] LibreOffice (soffice) not found")
            raise HTTPException(
                status_code=500,
                detail="PDF conversion isn't available on the server right now — the DOCX download works.",
            )
        except Exception as e:
            logger.exception("[build/pdf] conversion failed")
            raise HTTPException(
                status_code=500,
                detail=f"PDF conversion failed ({e}). The DOCX download works.",
            )
        with open(pdf_path, "rb") as f:
            data = f.read()
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe}.pdf"'},
    )


# ───────────────────────── AI Tailor (BETA) ─────────────────────────────────
# Separate from the deterministic /format flow: this uses an LLM to rewrite the
# resume for a job description and draft a cover letter + HR email. Factual data
# is re-stamped from the source in llm/tailor.py, so the model can't alter facts.

@app.get("/tailor/status")
def tailor_status():
    """Frontend checks this to show whether the AI feature is configured."""
    from llm.client import is_configured, get_model
    return {"configured": is_configured(), "model": get_model() if is_configured() else None}


@app.post("/tailor")
async def tailor_resume(
    file: UploadFile = File(None),
    plain_text: str = Form(None),
    job_description: str = Form(None),
    job_url: str = Form(None),
):
    """Parse the resume, gather the JD (pasted text and/or a scraped URL), and
    return three AI drafts for review. No files are written here."""
    from llm.client import LLMNotConfigured

    raw_text = await _read_resume_text(file, plain_text)

    jd_text = (job_description or "").strip()
    if (job_url or "").strip():
        from jd_scraper import fetch_jd
        scraped = fetch_jd(job_url)
        if scraped:
            jd_text = (jd_text + "\n\n" + scraped).strip() if jd_text else scraped
    if not jd_text:
        raise HTTPException(
            status_code=400,
            detail="Paste the job description (or a job link we can read) so we can tailor to it.",
        )

    # /tailor uses robust LLM extraction (handles any layout, then fact-validates
    # so it can't invent or lose an entry). /format keeps the deterministic parser.
    from llm.extract import extract_resume
    resume = extract_resume(raw_text)

    from llm.tailor import tailor
    try:
        result = tailor(resume, jd_text)
    except LLMNotConfigured as e:
        raise HTTPException(status_code=503, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=502, detail=str(e))
    except Exception as e:
        logger.exception("[tailor] generation failed")
        raise HTTPException(status_code=502, detail=f"AI tailoring failed: {e}. Please try again.")

    return JSONResponse(result)


def _render_tailor_docx(req: "TailorBuildRequest", path: str):
    """Render one of the three outputs as a DOCX at `path`."""
    if req.kind == "resume":
        _require_named_resume(req.tailored_resume)
        _render_docx(req.tailored_resume, path)
    elif req.kind == "cover_letter":
        from formatters.letter import format_cover_letter
        format_cover_letter(
            {
                "name": req.tailored_resume.get("name", ""),
                "headline": req.tailored_resume.get("headline", ""),
                "contact": req.tailored_resume.get("contact", {}),
                "cover_letter": req.cover_letter,
            },
            path,
        )
    elif req.kind == "email":
        from formatters.letter import format_email
        format_email(
            {
                "name": req.tailored_resume.get("name", ""),
                "contact": req.tailored_resume.get("contact", {}),
                "email": req.email,
            },
            path,
        )
    else:
        raise HTTPException(status_code=400, detail=f"Unknown output kind: {req.kind}")


def _tailor_filename(req: "TailorBuildRequest") -> str:
    base = _safe_filename(req.tailored_resume.get("name") or "Application")
    suffix = {"resume": "Resume", "cover_letter": "CoverLetter", "email": "Email"}.get(req.kind, "Document")
    return f"{base}_{suffix}"


@app.post("/tailor/build")
def tailor_build(req: TailorBuildRequest):
    """Render one approved AI output (resume/cover_letter/email) as DOCX or PDF
    and stream it back in one request — same single-request pattern as /build."""
    name = _tailor_filename(req)
    fmt = (req.format or "docx").lower()
    tmpdir = tempfile.mkdtemp(prefix="tailor-")
    try:
        docx_path = os.path.join(tmpdir, f"{name}.docx")
        _render_tailor_docx(req, docx_path)
        if fmt == "pdf":
            pdf_path = os.path.join(tmpdir, f"{name}.pdf")
            try:
                _convert_to_pdf(docx_path, pdf_path, tmpdir)
            except FileNotFoundError:
                raise HTTPException(
                    status_code=500,
                    detail="PDF conversion isn't available on the server right now — the DOCX download works.",
                )
            except Exception as e:
                logger.exception("[tailor/build] pdf conversion failed")
                raise HTTPException(status_code=500, detail=f"PDF conversion failed ({e}). The DOCX download works.")
            with open(pdf_path, "rb") as f:
                data = f.read()
            return Response(
                content=data,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{name}.pdf"'},
            )
        with open(docx_path, "rb") as f:
            data = f.read()
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("[tailor/build] could not generate document")
        raise HTTPException(status_code=500, detail=f"Couldn't generate the document: {e}")
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
    return Response(
        content=data,
        media_type=DOCX_MEDIA_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{name}.docx"'},
    )
