from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import logging
import re

logger = logging.getLogger("resume.formatter")

# The "List Bullet" style already draws a marker, so strip any leading bullet
# glyph the source carried (incl. Symbol/Wingdings PUA bullets like U+F0B7) to
# avoid double bullets — defence in depth on top of the structurer's stripping.
_LEADING_BULLET_RE = re.compile(
    r"^\s*[•·‣◦▪●∙○■◆➢➤»-]+\s*"
)


FONT_NAME = "Calibri"
NAME_SIZE = Pt(16)
TITLE_SIZE = Pt(10.5)
BODY_SIZE = Pt(9)
HEADER_SIZE = Pt(9.5)
MARGIN = Inches(0.20)
COBALT_BLUE = RGBColor(0, 71, 171)
HEADLINE_GRAY = RGBColor(80, 80, 80)

# US Letter page (8.5 in wide) — the US-market standard this product targets.
# Right-align tab stops sit at the page width minus the two side margins,
# expressed in twips (1 inch = 1440 twips).
LETTER_WIDTH = Inches(8.5)
LETTER_HEIGHT = Inches(11)
USABLE_WIDTH_TWIPS = int((8.5 - 2 * 0.20) * 1440)


def _set_spacing(paragraph, before=0, after=0, line_rule="auto", line=240):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:lineRule"), line_rule)
    spacing.set(qn("w:line"), str(line))
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def _set_margins(doc):
    section = doc.sections[0]
    section.page_width = LETTER_WIDTH
    section.page_height = LETTER_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN


def _add_name(doc, name: str):
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(name or "").upper())
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = NAME_SIZE
    run.font.color.rgb = COBALT_BLUE


def _add_headline(doc, headline: str):
    """The professional title line shown directly under the name (e.g.
    'Data Analyst | BI Analyst'). Subordinate to the name: smaller and muted."""
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(headline)
    run.font.name = FONT_NAME
    run.font.size = TITLE_SIZE
    run.font.color.rgb = HEADLINE_GRAY


def _add_link_styled_run(paragraph, text: str):
    """A plain run styled like a hyperlink (same cobalt blue + size) but with no
    target. Used when a URL is missing or invalid so the visible text is always
    preserved — a broken link never drops a word or aborts the document."""
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE
    run.font.color.rgb = COBALT_BLUE


def _add_hyperlink(paragraph, text: str, url: str):
    """Render `text` as an external hyperlink. If the URL is empty or python-docx
    can't build the relationship (e.g. a malformed/edited link), fall back to a
    styled plain run instead of crashing the whole document."""
    url = (url or "").strip()
    if not url:
        _add_link_styled_run(paragraph, text)
        return
    try:
        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
    except Exception:
        _add_link_styled_run(paragraph, text)
        return
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0047AB")
    rPr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(BODY_SIZE.pt * 2)))
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def _add_plain_run(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE


def _ensure_scheme(url: str) -> str:
    """Add https:// only when no scheme is present — never doubles it up."""
    url = url.strip()
    if url.lower().startswith(("http://", "https://", "mailto:", "tel:")):
        return url
    return "https://" + url.lstrip("/")


def _profile_url(raw: str, domain: str, handle_prefix: str) -> str:
    """Build a clickable target for a profile that may be a full URL, a bare
    domain path, or just a username/handle. The displayed text stays exactly
    as the user wrote it — this only affects where the hyperlink points."""
    raw = raw.strip()
    if not raw:
        return ""  # nothing to link — caller renders the label as plain text
    low = raw.lower()
    if low.startswith(("http://", "https://", "www.")) or domain in low:
        return _ensure_scheme(raw)
    # Bare username/handle — attach to the canonical profile path.
    return f"https://{handle_prefix}{raw.lstrip('/')}"


def _add_contact(doc, contact: dict):
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _clean(key):
        return (contact.get(key) or "").strip()

    segments = []
    phone = _clean("phone")
    if phone:
        segments.append(("plain", phone, None))
    email = _clean("email")
    if email:
        addr = email[len("mailto:"):] if email.lower().startswith("mailto:") else email
        href = email if email.lower().startswith("mailto:") else f"mailto:{email}"
        segments.append(("link", _clean("email_label") or addr, href))
    # LinkedIn / GitHub each have a display label (default "LinkedIn"/"GitHub")
    # separate from the link target, so the user can edit text and URL apart.
    linkedin, linkedin_label = _clean("linkedin"), _clean("linkedin_label")
    if linkedin or linkedin_label:
        url = _profile_url(linkedin, "linkedin.com", "www.linkedin.com/in/") if linkedin else ""
        segments.append(("link", linkedin_label or "LinkedIn", url))
    github, github_label = _clean("github"), _clean("github_label")
    if github or github_label:
        url = _profile_url(github, "github.com", "github.com/") if github else ""
        segments.append(("link", github_label or "GitHub", url))
    for link in contact.get("links") or []:
        url = (link.get("url") or "").strip()
        label = (link.get("label") or "").strip()
        if not url and not label:
            continue  # an empty link row carries nothing — skip it
        segments.append(("link", label or url, _ensure_scheme(url) if url else ""))
    location = _clean("location")
    if location:
        segments.append(("plain", location, None))

    # Every segment now has non-empty display text, so the "  |  " separator can
    # never dangle on an empty field.
    for i, (kind, text, url) in enumerate(segments):
        if i > 0:
            _add_plain_run(p, "  |  ")
        if kind == "link":
            _add_hyperlink(p, text, url)
        else:
            _add_plain_run(p, text)


def _add_section_header(doc, title: str):
    p = doc.add_paragraph()
    _set_spacing(p, before=4, after=1)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = HEADER_SIZE
    run.font.color.rgb = COBALT_BLUE

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_job_header(doc, title: str, company: str, location: str, start: str, end: str):
    p = doc.add_paragraph()
    _set_spacing(p, before=2, after=0)

    # Join only the parts that are actually present — an empty title/company
    # must never leave a dangling "  |  " before the next field.
    left_text = "  |  ".join(p.strip() for p in (title, company, location) if p and p.strip())
    date_text = " – ".join(d.strip() for d in (start, end) if d and d.strip())

    run_left = p.add_run(left_text)
    run_left.bold = True
    run_left.font.name = FONT_NAME
    run_left.font.size = BODY_SIZE

    # Right-align the date using a tab stop at page width
    tab = OxmlElement("w:tab")
    run_left._r.append(tab)

    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_stop = OxmlElement("w:tab")
    usable_width = USABLE_WIDTH_TWIPS
    tab_stop.set(qn("w:val"), "right")
    tab_stop.set(qn("w:pos"), str(usable_width))
    tabs.append(tab_stop)
    pPr.append(tabs)

    run_date = p.add_run(date_text)
    run_date.bold = True
    run_date.font.name = FONT_NAME
    run_date.font.size = BODY_SIZE


def _add_bullet(doc, text: str):
    text = _LEADING_BULLET_RE.sub("", (text or ""), count=1)
    if not text.strip():
        return  # skip an empty/whitespace-only bullet — no data to render
    p = doc.add_paragraph(style="List Bullet")
    _set_spacing(p, before=0, after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE

    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "180")
    ind.set(qn("w:hanging"), "180")
    existing = pPr.find(qn("w:ind"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(ind)


def _add_skills(doc, skills: dict):
    for category, items in skills.items():
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=1)
        run_label = p.add_run(f"{category}: ")
        run_label.bold = True
        run_label.font.name = FONT_NAME
        run_label.font.size = BODY_SIZE
        run_items = p.add_run(", ".join(items))
        run_items.font.name = FONT_NAME
        run_items.font.size = BODY_SIZE


def _add_prose(doc, text: str):
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE


def _add_inline_items(doc, items: list):
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=1)
    run = p.add_run(", ".join(items))
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE


def _add_additional_sections(doc, sections: list):
    """Render any non-standard resume heading captured by the parser.

    Each section carries a free-form heading plus a `style` chosen by the
    parser so we can format it like the closest standard section:
      - "skills" -> inline comma-joined line (e.g. Languages, Tools)
      - "list"   -> justified bullets
      - "prose"  -> justified paragraph(s)
    Nothing is dropped: unknown/blank styles fall back to whatever content exists.
    """
    for section in sections:
        heading = (section.get("heading") or "").strip()
        if not heading:
            continue

        items = section.get("items") or []
        text = section.get("text")
        style = (section.get("style") or "").strip().lower()

        # Don't emit an empty header if there's genuinely no content.
        if not items and not (text and text.strip()):
            continue

        _add_section_header(doc, heading)

        if style == "prose" and text and text.strip():
            _add_prose(doc, text.strip())
        elif style == "skills" and items:
            _add_inline_items(doc, items)
        elif style == "list" and items:
            for item in items:
                _add_bullet(doc, item)
        else:
            # Fallback so no information is ever lost, even on an odd style.
            if text and text.strip():
                _add_prose(doc, text.strip())
            for item in items:
                _add_bullet(doc, item)


# ─────────────────────────── per-section renderers ──────────────────────────
# Each renderer is self-contained (no shared state) so sections can be emitted
# in any order. Name/headline/contact/summary stay pinned at the top; everything
# below the summary is reorderable via `section_order`.

def _render_summary(doc, data):
    if data.get("summary"):
        _add_section_header(doc, "Professional Summary")
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(data["summary"])
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE


def _render_skills(doc, data):
    if data.get("skills"):
        _add_section_header(doc, "Technical Skills")
        _add_skills(doc, data["skills"])


def _render_experience(doc, data):
    if data.get("experience"):
        _add_section_header(doc, "Professional Experience")
        for job in data["experience"]:
            _add_job_header(
                doc,
                job.get("title", ""),
                job.get("company", ""),
                job.get("location", ""),
                job.get("start_date", ""),
                job.get("end_date", ""),
            )
            for bullet in job.get("bullets", []):
                _add_bullet(doc, bullet)


def _render_projects(doc, data):
    if data.get("projects"):
        _add_section_header(doc, "Projects")
        for project in data["projects"]:
            p = doc.add_paragraph()
            _set_spacing(p, before=2, after=0)
            run_name = p.add_run(project.get("name", ""))
            run_name.bold = True
            run_name.font.name = FONT_NAME
            run_name.font.size = BODY_SIZE
            if project.get("tech_stack"):
                run_stack = p.add_run(f"  |  {project['tech_stack']}")
                run_stack.font.name = FONT_NAME
                run_stack.font.size = BODY_SIZE
                run_stack.italic = True
            for bullet in project.get("bullets", []):
                _add_bullet(doc, bullet)


def _render_certifications(doc, data):
    if data.get("certifications"):
        _add_section_header(doc, "Certifications")
        for cert in data["certifications"]:
            p = doc.add_paragraph()
            _set_spacing(p, before=1, after=1)
            run_name = p.add_run(cert.get("name", ""))
            run_name.bold = True
            run_name.font.name = FONT_NAME
            run_name.font.size = BODY_SIZE
            if cert.get("issuer"):
                run_issuer = p.add_run(f"  |  {cert['issuer']}")
                run_issuer.font.name = FONT_NAME
                run_issuer.font.size = BODY_SIZE
            if cert.get("date"):
                tab = OxmlElement("w:tab")
                run_name._r.append(tab)
                pPr = p._p.get_or_add_pPr()
                tabs = OxmlElement("w:tabs")
                tab_stop = OxmlElement("w:tab")
                usable_width = USABLE_WIDTH_TWIPS
                tab_stop.set(qn("w:val"), "right")
                tab_stop.set(qn("w:pos"), str(usable_width))
                tabs.append(tab_stop)
                pPr.append(tabs)
                run_date = p.add_run(cert["date"])
                run_date.font.name = FONT_NAME
                run_date.font.size = BODY_SIZE
            for bullet in cert.get("bullets", []):
                _add_bullet(doc, bullet)


def _render_education(doc, data):
    if data.get("education"):
        _add_section_header(doc, "Education")
        for edu in data["education"]:
            p = doc.add_paragraph()
            _set_spacing(p, before=1, after=0)
            run_deg = p.add_run(edu.get("degree", ""))
            run_deg.bold = True
            run_deg.font.name = FONT_NAME
            run_deg.font.size = BODY_SIZE

            tab = OxmlElement("w:tab")
            run_deg._r.append(tab)

            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab_stop = OxmlElement("w:tab")
            usable_width = USABLE_WIDTH_TWIPS
            tab_stop.set(qn("w:val"), "right")
            tab_stop.set(qn("w:pos"), str(usable_width))
            tabs.append(tab_stop)
            pPr.append(tabs)

            if edu.get("graduation_date"):
                run_date = p.add_run(edu["graduation_date"])
                run_date.bold = True
                run_date.font.name = FONT_NAME
                run_date.font.size = BODY_SIZE

            # Keep BOTH institution and location even when one is missing — never
            # drop a field, and never leave a dangling "  |  ".
            inst_text = "  |  ".join(
                s for s in ((edu.get("institution") or "").strip(), (edu.get("location") or "").strip()) if s
            )
            p2 = doc.add_paragraph()
            _set_spacing(p2, before=0, after=1)
            run_inst = p2.add_run(inst_text)
            run_inst.font.name = FONT_NAME
            run_inst.font.size = BODY_SIZE

            # GPA on its own line (kept separate so the UI can hide it on request).
            if edu.get("gpa"):
                pg = doc.add_paragraph()
                _set_spacing(pg, before=0, after=1)
                run_gpa = pg.add_run(edu["gpa"])
                run_gpa.font.name = FONT_NAME
                run_gpa.font.size = BODY_SIZE

            # Coursework / honours / any other extra academic lines.
            for detail in edu.get("details") or []:
                _add_bullet(doc, detail)


def _render_additional_one(doc, data, index):
    sections = data.get("additional_sections") or []
    if 0 <= index < len(sections):
        _add_additional_sections(doc, [sections[index]])


_BODY_RENDERERS = {
    "skills": _render_skills,
    "experience": _render_experience,
    "projects": _render_projects,
    "certifications": _render_certifications,
    "education": _render_education,
}


def _default_body_order(n_additional: int) -> list:
    """Skills → Experience → Projects → Certifications → (additional) → Education."""
    additional = [f"additional-{i}" for i in range(n_additional)]
    return ["skills", "experience", "projects", "certifications", *additional, "education"]


def _resolve_section_order(data: dict) -> list:
    """The reorderable body order (everything after the summary). Honours a
    user-supplied ``section_order`` but always appends any present section the
    request omitted, so a section can never be silently dropped."""
    n_additional = len(data.get("additional_sections") or [])
    default = _default_body_order(n_additional)
    requested = data.get("section_order")
    if not (requested and isinstance(requested, list)):
        return default
    known = set(default)
    ordered = [k for k in requested if k in known]
    seen = set(ordered)
    for k in default:
        if k not in seen:
            ordered.append(k)
    return ordered


def _safe(label, fn, *args):
    """Run a render step but never let one bad field/section abort the document."""
    try:
        fn(*args)
    except Exception:
        logger.exception("[compact_ats] failed rendering %s", label)


def _render_body_section(doc, data, key):
    if key.startswith("additional-"):
        try:
            _render_additional_one(doc, data, int(key.split("-", 1)[1]))
        except Exception:
            logger.exception("[compact_ats] failed to render section '%s'", key)
        return
    renderer = _BODY_RENDERERS.get(key)
    if renderer:
        _safe(f"section '{key}'", renderer, doc, data)


def format_compact(data: dict, output_path: str):
    doc = Document()
    _set_margins(doc)

    # Remove default paragraph spacing globally
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE

    # Pinned header: name, headline, contact, summary. Each is isolated so a
    # single malformed field can never abort the whole document.
    _safe("name", _add_name, doc, data.get("name", ""))
    if data.get("headline"):
        _safe("headline", _add_headline, doc, data["headline"])
    if data.get("contact"):
        _safe("contact", _add_contact, doc, data["contact"])
    _safe("summary", _render_summary, doc, data)

    # Reorderable body (everything after the summary).
    for key in _resolve_section_order(data):
        _render_body_section(doc, data, key)

    doc.save(output_path)
    logger.info("[compact_ats] saved %s", output_path)
