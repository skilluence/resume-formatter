"""DOCX renderers for the AI Tailor cover letter and HR email.

Cover letter: a modern letterhead - left-aligned bold name, a cobalt role line, a
two-column icon contact block (phone | email / linkedin | location), then a
"COVER LETTER" heading with a cobalt rule, then the letter body. Cobalt accent
(#0047AB) matches the resume so the application reads as one set.

HR email: a lighter layout - bold "Subject:" line, greeting, body, signature, then
a slim cobalt contact footer (closer to a real email than a letterhead).

Body paragraphs honour **bold** keyword markup via _add_rich_runs. Business-letter
spacing (1" margins, 11pt Calibri).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from formatters.compact_ats import (
    FONT_NAME,
    COBALT_BLUE,
    _add_hyperlink,
    _add_plain_run,
    _add_rich_runs,
    _profile_url,
    _set_spacing,
)
from formatters.letter_icons import icon_png

LETTER_FONT_SIZE = Pt(11)
NAME_SIZE = Pt(17)
ROLE_SIZE = Pt(11.5)
NAME_DARK = RGBColor(0x22, 0x22, 0x22)
LETTER_MARGIN = Inches(1.0)
LETTER_WIDTH = Inches(8.5)
LETTER_HEIGHT = Inches(11)
COBALT_HEX = "0047AB"


def _setup(doc):
    section = doc.sections[0]
    section.page_width = LETTER_WIDTH
    section.page_height = LETTER_HEIGHT
    section.top_margin = LETTER_MARGIN
    section.bottom_margin = LETTER_MARGIN
    section.left_margin = LETTER_MARGIN
    section.right_margin = LETTER_MARGIN
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = LETTER_FONT_SIZE


def _clean(contact: dict, key: str) -> str:
    return (contact.get(key) or "").strip()


def _contact_items(contact: dict) -> list:
    """Ordered (icon, text, href) tuples for the present contact fields:
    phone, email, linkedin, location - laid out two per row to match the design."""
    contact = contact or {}
    items = []
    phone = _clean(contact, "phone")
    if phone:
        items.append(("phone", phone, None))
    email = _clean(contact, "email")
    if email:
        addr = email[len("mailto:"):] if email.lower().startswith("mailto:") else email
        href = email if email.lower().startswith("mailto:") else f"mailto:{email}"
        items.append(("email", _clean(contact, "email_label") or addr, href))
    linkedin, label = _clean(contact, "linkedin"), _clean(contact, "linkedin_label")
    if linkedin or label:
        url = _profile_url(linkedin, "linkedin.com", "www.linkedin.com/in/") if linkedin else ""
        items.append(("linkedin", label or "LinkedIn", url))
    location = _clean(contact, "location")
    if location:
        items.append(("location", location, None))
    return items


_ICON_LABEL = {"phone": "Tel:", "email": "Email:", "linkedin": "LinkedIn:", "location": "Location:"}


def _add_icon(paragraph, kind: str):
    """Inline cobalt icon, with a text-label fallback if the image can't be drawn."""
    run = paragraph.add_run()
    buf = icon_png(kind)
    if buf is not None:
        try:
            run.add_picture(buf, height=Pt(10))
            return
        except Exception:
            pass
    lbl = paragraph.add_run(_ICON_LABEL.get(kind, ""))
    lbl.bold = True
    lbl.font.name = FONT_NAME
    lbl.font.size = Pt(9.5)
    lbl.font.color.rgb = COBALT_BLUE


def _no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _add_contact_grid(doc, contact: dict):
    items = _contact_items(contact)
    if not items:
        return
    rows = (len(items) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    _no_table_borders(table)
    for i, (kind, text, href) in enumerate(items):
        cell = table.cell(i // 2, i % 2)
        para = cell.paragraphs[0]
        _set_spacing(para, before=1, after=1)
        _add_icon(para, kind)
        _add_plain_run(para, "  ")
        if href:
            _add_hyperlink(para, text, href)
        else:
            run = para.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)


def _bottom_border(paragraph, color=COBALT_HEX, sz="6", space="2"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _top_border(paragraph, color=COBALT_HEX, sz="6", space="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), space)
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def _add_letterhead(doc, name: str, headline: str, contact: dict):
    """Name (left, bold, near-black) -> role line (cobalt) -> 2-col icon contacts."""
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=1)
    run = p.add_run((name or "").strip())
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = NAME_SIZE
    run.font.color.rgb = NAME_DARK

    if (headline or "").strip():
        rp = doc.add_paragraph()
        _set_spacing(rp, before=0, after=8)
        rr = rp.add_run(headline.strip())
        rr.font.name = FONT_NAME
        rr.font.size = ROLE_SIZE
        rr.font.color.rgb = COBALT_BLUE

    _add_contact_grid(doc, contact)


def _add_heading(doc, title: str):
    p = doc.add_paragraph()
    _set_spacing(p, before=10, after=8)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = COBALT_BLUE
    _bottom_border(p)


def _body_paragraph(doc, text: str, after=8):
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=after)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_rich_runs(p, text, LETTER_FONT_SIZE)
    return p


def _signature(doc, name: str):
    sp = doc.add_paragraph()
    _set_spacing(sp, before=0, after=0)
    run = sp.add_run(name)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = LETTER_FONT_SIZE


def format_cover_letter(data: dict, output_path: str):
    """data: {name, headline, contact, cover_letter:{greeting, body_paragraphs,
    closing, signature}}."""
    doc = Document()
    _setup(doc)
    letter = data.get("cover_letter") or {}

    _add_letterhead(doc, data.get("name", ""), data.get("headline", ""), data.get("contact"))
    _add_heading(doc, "Cover Letter")

    if letter.get("greeting"):
        _body_paragraph(doc, letter["greeting"], after=8)
    for para in letter.get("body_paragraphs") or []:
        if str(para).strip():
            _body_paragraph(doc, str(para), after=8)
    if letter.get("closing"):
        _body_paragraph(doc, letter["closing"], after=0)
    if letter.get("signature"):
        _signature(doc, letter["signature"])

    doc.save(output_path)


def format_email(data: dict, output_path: str):
    """data: {name, contact, email:{subject, greeting, body_paragraphs, closing,
    signature}}. Light layout: a bold Subject line, the body, then a slim cobalt
    contact footer."""
    doc = Document()
    _setup(doc)
    email = data.get("email") or {}

    if email.get("subject"):
        sp = doc.add_paragraph()
        _set_spacing(sp, before=0, after=10)
        lbl = sp.add_run("Subject: ")
        lbl.bold = True
        lbl.font.name = FONT_NAME
        lbl.font.size = LETTER_FONT_SIZE
        lbl.font.color.rgb = COBALT_BLUE
        _add_rich_runs(sp, email["subject"], LETTER_FONT_SIZE)

    if email.get("greeting"):
        _body_paragraph(doc, email["greeting"], after=8)
    for para in email.get("body_paragraphs") or []:
        if str(para).strip():
            _body_paragraph(doc, str(para), after=8)
    if email.get("closing"):
        _body_paragraph(doc, email["closing"], after=0)
    if email.get("signature"):
        _signature(doc, email["signature"])

    _add_email_footer(doc, data.get("contact"))
    doc.save(output_path)


def _add_email_footer(doc, contact: dict):
    """A slim cobalt-ruled contact line under the signature (email | LinkedIn |
    location), with clickable email/LinkedIn - mirrors the resume's accent."""
    items = _contact_items(contact)
    items = [it for it in items if it[0] in ("email", "linkedin", "location", "phone")]
    if not items:
        return
    p = doc.add_paragraph()
    _set_spacing(p, before=10, after=0)
    _top_border(p)
    for i, (kind, text, href) in enumerate(items):
        if i > 0:
            _add_plain_run(p, "   |   ")
        if href:
            _add_hyperlink(p, text, href)
        else:
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = Pt(9.5)
            run.font.color.rgb = COBALT_BLUE
