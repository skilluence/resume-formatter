"""Cover letter / email DOCX renderers + the **bold** rich-run helper."""
import os
import tempfile

from docx import Document

from formatters.letter import format_cover_letter, format_email
from formatters.compact_ats import _BOLD_RE


DATA = {
    "name": "Jane Doe",
    "contact": {"email": "jane@example.com", "linkedin": "jane-doe", "location": "Austin, TX"},
    "cover_letter": {
        "greeting": "Dear Hiring Manager,",
        "body_paragraphs": ["I bring strong **SQL** and **Tableau** skills."],
        "closing": "Sincerely,",
        "signature": "Jane Doe",
    },
    "email": {
        "subject": "Application - Data Analyst",
        "greeting": "Dear Hiring Manager,",
        "body_paragraphs": ["My **resume** and **cover letter** are attached."],
        "closing": "Best regards,",
        "signature": "Jane Doe",
    },
}


def _hyperlink_targets(docx_path):
    doc = Document(docx_path)
    rels = doc.part.rels
    return [r.target_ref for r in rels.values() if "hyperlink" in r.reltype]


def test_cover_letter_renders_with_bold_and_links():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "cl.docx")
        format_cover_letter(DATA, path)
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        # Markup is stripped; the literal ** must not survive into the document.
        assert "**" not in text
        assert "SQL" in text and "Tableau" in text
        # A bold run exists for the keyword.
        bolded = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
        assert "SQL" in bolded
        # Email is a real hyperlink.
        assert any("mailto:jane@example.com" in t for t in _hyperlink_targets(path))


def test_email_renders_subject_and_links():
    with tempfile.TemporaryDirectory() as d:
        path = os.path.join(d, "email.docx")
        format_email(DATA, path)
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Subject:" in text
        assert "Application - Data Analyst" in text
        assert "**" not in text
        targets = _hyperlink_targets(path)
        assert any("linkedin.com" in t for t in targets)


def test_bold_regex_matches_double_asterisks():
    assert _BOLD_RE.findall("a **b** c **d**") == ["b", "d"]
    assert _BOLD_RE.findall("no markup here") == []
